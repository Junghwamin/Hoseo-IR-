"""
IR센터 앱 종합 E2E 테스트 스크립트

Streamlit 앱의 모든 핵심 모듈을 실제 데이터로 테스트한다.
각 단계(1~5)의 데이터 흐름을 시뮬레이션하여
실제 앱 실행 시 발생할 수 있는 문제를 사전에 탐지한다.
"""

import sys
import os
import traceback
from pathlib import Path
from io import BytesIO

# 프로젝트 루트를 sys.path에 추가
PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))
os.chdir(str(PROJECT_ROOT))

passed = 0
failed = 0
warnings = []
errors = []

def test(name):
    def decorator(func):
        global passed, failed
        try:
            result = func()
            if result is None or result is True:
                print(f"  ✅ PASS: {name}")
                passed += 1
            else:
                print(f"  ❌ FAIL: {name} - {result}")
                failed += 1
                errors.append(f"{name}: {result}")
        except Exception as e:
            print(f"  ❌ FAIL: {name} - {e}")
            traceback.print_exc()
            failed += 1
            errors.append(f"{name}: {e}")
    return decorator

def warn(msg):
    print(f"  ⚠️  WARN: {msg}")
    warnings.append(msg)


# ================================================================
print("\n" + "=" * 60)
print("  IR센터 앱 종합 E2E 테스트")
print("=" * 60)

# ================================================================
print("\n📦 [모듈 Import 테스트]")
# ================================================================

@test("config 모듈 import")
def _():
    from report_app.config import UNIVERSITY, COMPARE_GROUP, NATIONAL_CSV, REGIONAL_CSV, IS_CLOUD, REPORT_DIR
    assert UNIVERSITY == "호서대학교"
    assert len(COMPARE_GROUP) == 5
    assert "호서대학교" in COMPARE_GROUP

@test("data_loader 모듈 import")
def _():
    from report_app import data_loader as dl
    assert hasattr(dl, 'load_all_data')
    assert hasattr(dl, 'get_hoseo_trend')
    assert hasattr(dl, 'get_averages')
    assert hasattr(dl, 'get_rank_changes')
    assert hasattr(dl, 'get_yoy_changes')
    assert hasattr(dl, 'get_compare_group_data')
    assert hasattr(dl, 'get_available_years')

@test("chart_generator 모듈 import")
def _():
    from report_app import chart_generator as cg
    assert hasattr(cg, 'create_trend_chart')
    assert hasattr(cg, 'create_comparison_bar')
    assert hasattr(cg, 'create_avg_comparison')
    assert hasattr(cg, 'create_rank_trend_chart')
    assert hasattr(cg, 'create_compare_group_bar')

@test("gpt_reporter 모듈 import")
def _():
    from report_app import gpt_reporter as gpt
    assert hasattr(gpt, 'generate_trend_narrative')
    assert hasattr(gpt, 'generate_comparison_narrative')
    assert hasattr(gpt, 'generate_regional_narrative')
    assert hasattr(gpt, 'generate_yoy_narrative')

@test("report_builder 모듈 import")
def _():
    from report_app import report_builder as rb
    assert hasattr(rb, 'build_report')

@test("전처리 스크립트 import")
def _():
    import importlib.util
    preprocess_path = PROJECT_ROOT / "전임교원_연구실적_전처리.py"
    assert preprocess_path.exists(), f"전처리 스크립트 없음: {preprocess_path}"
    spec = importlib.util.spec_from_file_location("preprocessor", str(preprocess_path))
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    assert hasattr(mod, 'main')
    assert hasattr(mod, 'process_in_memory')


# ================================================================
print("\n📂 [데이터 파일 검증]")
# ================================================================

@test("output/ 디렉터리 존재")
def _():
    assert (PROJECT_ROOT / "output").is_dir()

@test("전체_대학_데이터.csv 존재 및 로드")
def _():
    import pandas as pd
    csv_path = PROJECT_ROOT / "output" / "전체_대학_데이터.csv"
    assert csv_path.exists(), f"파일 없음: {csv_path}"
    df = pd.read_csv(csv_path, encoding="utf-8-sig")
    assert len(df) > 0, "빈 DataFrame"
    required_cols = ["연도", "학교명", "전임교원수", "SCI/SCOPUS논문수", "1인당논문수", "전국순위"]
    for col in required_cols:
        assert col in df.columns, f"컬럼 누락: {col}"
    print(f"       → {len(df)}행, 연도: {sorted(df['연도'].unique())}")

@test("충청권_순위.csv 존재 및 로드")
def _():
    import pandas as pd
    csv_path = PROJECT_ROOT / "output" / "충청권_순위.csv"
    assert csv_path.exists(), f"파일 없음: {csv_path}"
    df = pd.read_csv(csv_path, encoding="utf-8-sig")
    assert len(df) > 0, "빈 DataFrame"
    required_cols = ["연도", "학교명", "전임교원수", "SCI/SCOPUS논문수", "1인당논문수", "충청권순위", "전국순위"]
    for col in required_cols:
        assert col in df.columns, f"컬럼 누락: {col}"
    print(f"       → {len(df)}행, 연도: {sorted(df['연도'].unique())}")

@test("config 디렉터리 존재")
def _():
    assert (PROJECT_ROOT / "config").is_dir()
    assert (PROJECT_ROOT / "config" / "universities.json").exists()
    assert (PROJECT_ROOT / "config" / "regions.json").exists()


# ================================================================
print("\n📊 [Step 1: 데이터 로드 테스트]")
# ================================================================

import pandas as pd
from report_app import data_loader as dl
from report_app.config import NATIONAL_CSV, REGIONAL_CSV

@test("load_all_data() - CSV에서 로드")
def _():
    nat_df, reg_df = dl.load_all_data()
    assert len(nat_df) > 0
    assert len(reg_df) > 0
    return True

nat_df, reg_df = dl.load_all_data()
years = dl.get_available_years(nat_df)
print(f"       → 사용 가능 연도: {years}")
latest_year = max(years)
print(f"       → 최신 연도: {latest_year}")

@test("load_all_data() - DataFrame 직접 전달")
def _():
    nat2, reg2 = dl.load_all_data(national_df=nat_df, regional_df=reg_df)
    assert nat2.equals(nat_df)
    assert reg2.equals(reg_df)

@test("호서대 데이터 존재 확인")
def _():
    hoseo_nat = nat_df[nat_df["학교명"] == "호서대학교"]
    hoseo_reg = reg_df[reg_df["학교명"] == "호서대학교"]
    assert len(hoseo_nat) > 0, "전국 데이터에 호서대 없음"
    assert len(hoseo_reg) > 0, "충청권 데이터에 호서대 없음"
    print(f"       → 전국 {len(hoseo_nat)}행, 충청권 {len(hoseo_reg)}행")


# ================================================================
print("\n📋 [Step 2: 통계 계산 테스트]")
# ================================================================

@test("get_hoseo_trend()")
def _():
    trend = dl.get_hoseo_trend(nat_df, reg_df)
    assert len(trend) > 0, "빈 딕셔너리"
    for year, data in trend.items():
        assert "논문수" in data, f"{year}년: 논문수 키 없음"
        assert "전임교원수" in data
        assert "1인당논문수" in data
        assert "전국순위" in data
        assert data["전임교원수"] > 0, f"{year}년: 전임교원수가 0"
    print(f"       → {len(trend)}개년 데이터")
    for y, d in sorted(trend.items()):
        print(f"         {y}년: 1인당논문수={d['1인당논문수']}, 충청권={d.get('충청권순위')}, 전국={d['전국순위']}")

hoseo_trend = dl.get_hoseo_trend(nat_df, reg_df)

@test("get_averages()")
def _():
    avgs = dl.get_averages(nat_df, reg_df)
    assert len(avgs) > 0
    for year, data in avgs.items():
        assert "전국평균" in data
        assert "충청권평균" in data
        assert "비교군평균" in data
        assert data["전국평균"] > 0, f"{year}년: 전국평균이 0"
    print(f"       → {latest_year}년 전국평균={avgs[latest_year]['전국평균']}")

averages = dl.get_averages(nat_df, reg_df)

@test("get_rank_changes()")
def _():
    ranks = dl.get_rank_changes(nat_df, reg_df)
    assert len(ranks) > 0
    for year, data in ranks.items():
        assert "충청권순위" in data
        assert "전국순위" in data
    # 최소 2년 이상이면 변화값 존재
    if len(ranks) >= 2:
        years_sorted = sorted(ranks.keys())
        second_year = years_sorted[1]
        assert ranks[second_year].get("전국순위_변화") is not None, "순위 변화 계산 안 됨"

rank_changes = dl.get_rank_changes(nat_df, reg_df)

@test("get_yoy_changes()")
def _():
    yoy = dl.get_yoy_changes(reg_df, latest_year)
    assert "상위" in yoy
    assert "하위" in yoy
    assert "호서" in yoy
    if yoy["호서"]:
        assert "증감률" in yoy["호서"]
        print(f"       → 호서대 증감률: {yoy['호서']['증감률']:+.1f}%")
    else:
        warn(f"호서대 YoY 데이터 없음 ({latest_year}년)")

yoy_changes = dl.get_yoy_changes(reg_df, latest_year)

@test("get_compare_group_data()")
def _():
    cmpd = dl.get_compare_group_data(nat_df, reg_df, latest_year)
    assert len(cmpd) > 0, "비교군 데이터 없음"
    hoseo_found = any(d["학교명"] == "호서대학교" for d in cmpd)
    assert hoseo_found, "비교군에 호서대 없음"
    print(f"       → {len(cmpd)}개교 데이터")
    for d in cmpd:
        print(f"         {d['학교명']}: 1인당={d['1인당논문수']}, 전국={d['전국순위']}")

compare_data = dl.get_compare_group_data(nat_df, reg_df, latest_year)


# ================================================================
print("\n📈 [Step 3: 차트 생성 테스트]")
# ================================================================

from report_app import chart_generator as cg

@test("create_trend_chart() - 연도별 추이 라인차트")
def _():
    buf = cg.create_trend_chart(hoseo_trend, averages)
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    data = buf.read()
    assert len(data) > 1000, f"이미지가 너무 작음: {len(data)} bytes"
    # PNG magic number
    buf.seek(0)
    assert buf.read(4) == b'\x89PNG', "PNG 포맷이 아님"
    print(f"       → {len(data):,} bytes")

@test("create_comparison_bar() - 충청권 막대차트")
def _():
    buf = cg.create_comparison_bar(reg_df, latest_year)
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    data = buf.read()
    assert len(data) > 1000
    print(f"       → {len(data):,} bytes")

@test("create_avg_comparison() - 평균 비교 바차트")
def _():
    buf = cg.create_avg_comparison(hoseo_trend, averages, latest_year)
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    data = buf.read()
    assert len(data) > 1000
    print(f"       → {len(data):,} bytes")

@test("create_rank_trend_chart() - 순위 변화 라인차트")
def _():
    buf = cg.create_rank_trend_chart(rank_changes)
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    data = buf.read()
    assert len(data) > 1000
    print(f"       → {len(data):,} bytes")

@test("create_compare_group_bar() - 비교군 막대차트")
def _():
    buf = cg.create_compare_group_bar(compare_data, latest_year)
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    data = buf.read()
    assert len(data) > 1000
    print(f"       → {len(data):,} bytes")

# 차트 생성 (Step 5 테스트용)
charts = {
    "trend":   cg.create_trend_chart(hoseo_trend, averages),
    "bar":     cg.create_comparison_bar(reg_df, latest_year),
    "avg":     cg.create_avg_comparison(hoseo_trend, averages, latest_year),
    "rank":    cg.create_rank_trend_chart(rank_changes),
    "compare": cg.create_compare_group_bar(compare_data, latest_year),
}


# ================================================================
print("\n🤖 [Step 4: GPT Reporter 구조 테스트 (API 호출 없음)]")
# ================================================================

from report_app import gpt_reporter as gpt
from report_app.config import GPT_MODEL, GPT_MAX_TOKENS, GPT_TEMPERATURE

@test("GPT 설정값 확인")
def _():
    assert GPT_MODEL in ["gpt-4o", "gpt-4", "gpt-3.5-turbo"], f"알 수 없는 모델: {GPT_MODEL}"
    assert 500 <= GPT_MAX_TOKENS <= 4000
    assert 0 <= GPT_TEMPERATURE <= 1
    print(f"       → model={GPT_MODEL}, max_tokens={GPT_MAX_TOKENS}, temp={GPT_TEMPERATURE}")

@test("시스템 프롬프트 구조 확인")
def _():
    assert hasattr(gpt, '_SYSTEM_PROMPT')
    prompt = gpt._SYSTEM_PROMPT
    assert "대학" in prompt, "프롬프트에 '대학' 키워드 없음"
    assert "불릿" in prompt or "•" in prompt, "프롬프트에 불릿 포인트 지시 없음"
    assert len(prompt) > 50, "프롬프트가 너무 짧음"

@test("_call_gpt 함수 시그니처 확인")
def _():
    import inspect
    sig = inspect.signature(gpt._call_gpt)
    params = list(sig.parameters.keys())
    assert "client" in params
    assert "user_content" in params


# ================================================================
print("\n📄 [Step 5: Word 보고서 생성 테스트]")
# ================================================================

from report_app import report_builder as rb

@test("build_report() - GPT 서술 없이 보고서 생성")
def _():
    narratives = {"trend": "", "comparison": "", "regional": "", "yoy": ""}
    buf = rb.build_report(
        year=latest_year,
        hoseo_trend=hoseo_trend,
        averages=averages,
        compare_data=compare_data,
        yoy_changes=yoy_changes,
        rank_changes=rank_changes,
        charts=charts,
        narratives=narratives,
    )
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    data = buf.read()
    assert len(data) > 5000, f"보고서가 너무 작음: {len(data)} bytes"
    # DOCX magic number (PK ZIP)
    buf.seek(0)
    assert buf.read(2) == b'PK', "DOCX(ZIP) 포맷이 아님"
    print(f"       → {len(data):,} bytes (서술 없음)")

@test("build_report() - 더미 GPT 서술 포함 보고서 생성")
def _():
    narratives = {
        "trend": "• 호서대학교의 1인당 논문수는 꾸준히 증가하고 있습니다.\n• 전국 평균 대비 상회하는 성과를 보여줍니다.",
        "comparison": "• 비교군 5개교 중 상위권에 위치합니다.",
        "regional": "• 충청권 내 순위가 상승 추세입니다.",
        "yoy": "• 전년 대비 증가율이 양호합니다.",
    }
    buf = rb.build_report(
        year=latest_year,
        hoseo_trend=hoseo_trend,
        averages=averages,
        compare_data=compare_data,
        yoy_changes=yoy_changes,
        rank_changes=rank_changes,
        charts=charts,
        narratives=narratives,
    )
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    data = buf.read()
    assert len(data) > 10000, f"보고서가 너무 작음: {len(data)} bytes"
    print(f"       → {len(data):,} bytes (서술 포함)")

@test("build_report() - 보고서 내용 검증 (python-docx로 열기)")
def _():
    from docx import Document
    narratives = {
        "trend": "• 테스트 서술입니다.",
        "comparison": "", "regional": "", "yoy": "",
    }
    buf = rb.build_report(
        year=latest_year,
        hoseo_trend=hoseo_trend,
        averages=averages,
        compare_data=compare_data,
        yoy_changes=yoy_changes,
        rank_changes=rank_changes,
        charts=charts,
        narratives=narratives,
    )
    buf.seek(0)
    doc = Document(buf)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "호서대학교" in full_text, "보고서에 '호서대학교' 없음"
    assert str(latest_year) in full_text, f"보고서에 '{latest_year}' 없음"
    assert "테스트 서술" in full_text, "보고서에 서술 텍스트 미반영"
    # 표 확인
    assert len(doc.tables) > 0, "보고서에 표가 없음"
    print(f"       → 단락 {len(doc.paragraphs)}개, 표 {len(doc.tables)}개")


# ================================================================
print("\n🔧 [전처리 스크립트 테스트 (process_in_memory)]")
# ================================================================

@test("process_in_memory() - Raw 파일로 테스트")
def _():
    raw_dir = PROJECT_ROOT / "Raw data"
    if not raw_dir.exists():
        warn("Raw data/ 폴더 없음 - process_in_memory 테스트 스킵")
        return True

    xlsx_files = list(raw_dir.glob("*.xlsx"))
    if not xlsx_files:
        warn("Raw data/ 에 xlsx 파일 없음 - process_in_memory 테스트 스킵")
        return True

    import importlib.util
    spec = importlib.util.spec_from_file_location(
        "preprocessor", str(PROJECT_ROOT / "전임교원_연구실적_전처리.py")
    )
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)

    uploaded = {}
    for f in xlsx_files:
        uploaded[f.name] = f.read_bytes()

    nat_mem, reg_mem = mod.process_in_memory(uploaded)
    assert len(nat_mem) > 0, "전국 데이터 빈 DataFrame"
    assert len(reg_mem) > 0, "충청권 데이터 빈 DataFrame"

    # 컬럼 구조 확인
    for col in ["연도", "학교명", "전임교원수", "SCI/SCOPUS논문수", "1인당논문수", "전국순위"]:
        assert col in nat_mem.columns, f"전국 데이터 컬럼 누락: {col}"
    for col in ["연도", "학교명", "전임교원수", "SCI/SCOPUS논문수", "1인당논문수", "충청권순위", "전국순위"]:
        assert col in reg_mem.columns, f"충청권 데이터 컬럼 누락: {col}"

    print(f"       → 전국 {len(nat_mem)}행, 충청권 {len(reg_mem)}행")


# ================================================================
print("\n🌐 [Streamlit Cloud 호환성 테스트]")
# ================================================================

@test("IS_CLOUD 플래그 정상 동작 (로컬=False)")
def _():
    from report_app.config import IS_CLOUD
    assert IS_CLOUD is False or IS_CLOUD == False, f"로컬인데 IS_CLOUD={IS_CLOUD}"

@test("packages.txt 존재 (Cloud 폰트 설치용)")
def _():
    pkg_path = PROJECT_ROOT / "packages.txt"
    assert pkg_path.exists(), "packages.txt 없음"
    content = pkg_path.read_text().strip()
    assert "fonts-nanum" in content, "packages.txt에 fonts-nanum 없음"

@test("requirements.txt 존재 및 필수 패키지 포함")
def _():
    req_path = PROJECT_ROOT / "requirements.txt"
    assert req_path.exists(), "requirements.txt 없음"
    content = req_path.read_text()
    required = ["streamlit", "pandas", "matplotlib", "openai", "python-docx", "python-dotenv", "openpyxl"]
    missing = [pkg for pkg in required if pkg not in content.lower()]
    if missing:
        return f"requirements.txt에 누락: {missing}"

@test(".streamlit/secrets.toml 템플릿 존재")
def _():
    secrets_path = PROJECT_ROOT / ".streamlit" / "secrets.toml"
    if not secrets_path.exists():
        warn(".streamlit/secrets.toml 없음 (Cloud에서는 대시보드에서 설정)")
        return True

@test(".gitignore에 secrets.toml 포함")
def _():
    gitignore = PROJECT_ROOT / ".gitignore"
    if gitignore.exists():
        content = gitignore.read_text()
        if "secrets.toml" not in content:
            warn(".gitignore에 secrets.toml 미포함 - 시크릿 유출 위험")
    return True

@test("app.py sys.path 설정 확인")
def _():
    app_path = PROJECT_ROOT / "report_app" / "app.py"
    content = app_path.read_text(encoding="utf-8")
    assert "sys.path.insert" in content, "sys.path 설정 누락 - Cloud에서 ModuleNotFoundError 발생"

@test("chart_generator.py Linux 폰트 등록 확인")
def _():
    cg_path = PROJECT_ROOT / "report_app" / "chart_generator.py"
    content = cg_path.read_text(encoding="utf-8")
    assert "NanumGothic" in content, "NanumGothic 폰트 설정 누락"
    assert "addfont" in content, "addfont() 호출 누락 - Cloud에서 폰트 미등록"


# ================================================================
print("\n🔍 [앱 코드 정적 분석 - 잠재적 버그 탐지]")
# ================================================================

@test("app.py - text_area에 value= 미사용 확인 (key 충돌 방지)")
def _():
    app_path = PROJECT_ROOT / "report_app" / "app.py"
    content = app_path.read_text(encoding="utf-8")
    # text_area에 value=와 key=가 동시에 있으면 문제
    import re
    text_area_calls = re.findall(r'st\.text_area\([^)]+\)', content, re.DOTALL)
    for call in text_area_calls:
        if 'value=' in call and 'key=' in call:
            return f"text_area에 value=와 key= 동시 사용 발견: {call[:80]}..."
    return True

@test("app.py - _go() 함수에서 narrative 백업/복원 확인")
def _():
    app_path = PROJECT_ROOT / "report_app" / "app.py"
    content = app_path.read_text(encoding="utf-8")
    assert "_saved_narrative_trend" in content, "_saved_ 백업 키 누락"
    assert "_saved_narrative_comparison" in content
    assert "_saved_narrative_regional" in content
    assert "_saved_narrative_yoy" in content

@test("app.py - Step 5에서 _saved_ 키 우선 참조 확인")
def _():
    app_path = PROJECT_ROOT / "report_app" / "app.py"
    content = app_path.read_text(encoding="utf-8")
    # Step 5의 narratives 딕셔너리에서 _saved_ 키를 먼저 확인하는지
    assert "_saved_narrative_trend" in content
    # .get("_saved_narrative_trend", ...) 패턴 확인
    assert 'get("_saved_narrative_trend"' in content, "Step 5에서 _saved_ 우선 참조 없음"

@test("app.py - _get_api_key() 플레이스홀더 필터링")
def _():
    app_path = PROJECT_ROOT / "report_app" / "app.py"
    content = app_path.read_text(encoding="utf-8")
    assert '여기에' in content, "_get_api_key()에 플레이스홀더 필터 없음"

@test("data_loader.py - 0으로 나누기 방지")
def _():
    dl_path = PROJECT_ROOT / "report_app" / "data_loader.py"
    content = dl_path.read_text(encoding="utf-8")
    # get_yoy_changes에서 0으로 나누기 방지 확인
    assert "1인당논문수_이전" in content
    assert "> 0" in content, "0으로 나누기 방지 로직 없음"

@test("report_builder.py - 빈 narratives 처리")
def _():
    rb_path = PROJECT_ROOT / "report_app" / "report_builder.py"
    content = rb_path.read_text(encoding="utf-8")
    # narratives.get()으로 안전하게 접근하는지
    assert 'narratives.get(' in content, "narratives 안전 접근 없음"

@test("전처리 스크립트 - config 경로 확인")
def _():
    pp_path = PROJECT_ROOT / "전임교원_연구실적_전처리.py"
    content = pp_path.read_text(encoding="utf-8")
    # process_in_memory에서 config 경로가 __file__ 기준인지
    assert 'Path(__file__).parent' in content, "config 경로가 __file__ 기준이 아님"


# ================================================================
print("\n🔄 [엣지 케이스 테스트]")
# ================================================================

@test("연도 1개만 있을 때 YoY 처리")
def _():
    single_year_reg = reg_df[reg_df["연도"] == latest_year].copy()
    yoy = dl.get_yoy_changes(single_year_reg, latest_year)
    # 전년도 데이터 없으므로 빈 결과여야 함 (crash 안 되면 OK)
    assert isinstance(yoy, dict)
    assert "상위" in yoy
    assert "하위" in yoy

@test("빈 charts 딕셔너리로 보고서 생성")
def _():
    empty_charts = {"trend": None, "bar": None, "avg": None, "rank": None, "compare": None}
    narratives = {"trend": "", "comparison": "", "regional": "", "yoy": ""}
    buf = rb.build_report(
        year=latest_year,
        hoseo_trend=hoseo_trend,
        averages=averages,
        compare_data=compare_data,
        yoy_changes=yoy_changes,
        rank_changes=rank_changes,
        charts=empty_charts,
        narratives=narratives,
    )
    assert isinstance(buf, BytesIO)
    buf.seek(0)
    assert len(buf.read()) > 1000

@test("비교군에 없는 대학 처리")
def _():
    # 비교군에 존재하지 않는 대학이 있을 때
    cmpd = dl.get_compare_group_data(nat_df, reg_df, latest_year)
    # 반환된 데이터의 학교명이 비교군 목록에 있는지
    from report_app.config import COMPARE_GROUP
    for d in cmpd:
        assert d["학교명"] in COMPARE_GROUP, f"{d['학교명']}이 비교군 목록에 없음"

@test("차트 BytesIO seek(0) 후 재사용 가능")
def _():
    buf = cg.create_trend_chart(hoseo_trend, averages)
    # 첫 번째 읽기
    buf.seek(0)
    data1 = buf.read()
    # 두 번째 읽기 (보고서 삽입 시 재사용)
    buf.seek(0)
    data2 = buf.read()
    assert data1 == data2, "BytesIO 재사용 불가"
    assert len(data1) > 0


# ================================================================
print("\n🌐 [HTTP 접근 테스트]")
# ================================================================

import urllib.request

@test("Streamlit 앱 HTTP 200 응답")
def _():
    try:
        resp = urllib.request.urlopen("http://localhost:8502", timeout=10)
        assert resp.status == 200
        html = resp.read().decode()
        assert len(html) > 100
        print(f"       → {len(html):,} bytes HTML")
    except Exception as e:
        return f"HTTP 접근 실패: {e}"

@test("Streamlit health endpoint")
def _():
    try:
        resp = urllib.request.urlopen("http://localhost:8502/_stcore/health", timeout=5)
        assert resp.status == 200
    except Exception as e:
        warn(f"health endpoint 실패 (무시 가능): {e}")
        return True


# ================================================================
# 최종 결과
# ================================================================
print("\n" + "=" * 60)
print("  E2E 테스트 결과 요약")
print("=" * 60)
print(f"\n  ✅ PASS: {passed}")
print(f"  ❌ FAIL: {failed}")
print(f"  ⚠️  WARN: {len(warnings)}")

if warnings:
    print(f"\n  경고 사항:")
    for w in warnings:
        print(f"    ⚠️  {w}")

if errors:
    print(f"\n  실패 항목:")
    for e in errors:
        print(f"    ❌ {e}")

print("\n" + "=" * 60)

if failed > 0:
    sys.exit(1)
else:
    print("  🎉 모든 테스트 통과!")
    sys.exit(0)
