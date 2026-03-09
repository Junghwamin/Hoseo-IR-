"""
Word (.docx) 보고서 생성 모듈 (python-docx)

build_report() 함수가 모든 섹션을 조립하여 BytesIO Word 파일을 반환한다.
"""

from __future__ import annotations

from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from report_app.config import (
    COMPARE_GROUP_NAME,
    REPORT_FONT,
    REPORT_TITLE,
    UNIVERSITY,
)


# ---------------------------------------------------------------------------
# 헬퍼: 스타일 적용
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """셀 배경색 설정 (hex: 'RRGGBB')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_font(run, size_pt: int = 10, bold: bool = False, color_hex: str | None = None):
    """Run 폰트 스타일 설정."""
    run.font.name = REPORT_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_heading(doc: Document, text: str, level: int = 1):
    """섹션 제목 추가."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_font(run, size_pt=13 if level == 1 else 11, bold=True, color_hex="1f4e79")


def _add_narrative(doc: Document, text: str):
    """GPT 서술 텍스트를 불릿 단락으로 추가."""
    for line in text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        # 불릿 기호 제거 후 단락 추가
        clean = line.lstrip("•-· ").strip()
        if not clean:
            continue
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(clean)
        _set_font(run, size_pt=10)
        p.paragraph_format.space_after = Pt(3)


def _add_image(doc: Document, img_buf: BytesIO, width_cm: float = 15.0):
    """BytesIO 이미지를 Word에 삽입한다."""
    img_buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_buf, width=Cm(width_cm))
    doc.add_paragraph()  # 이미지 후 여백


# ---------------------------------------------------------------------------
# 표 생성 헬퍼
# ---------------------------------------------------------------------------

def _add_trend_table(doc: Document, hoseo_trend: dict[int, dict], region_name: str = "충청권"):
    """연도별 호서대 수치 표."""
    years = sorted(hoseo_trend.keys())
    headers = ["연도", "전임교원수", "SCI/SCOPUS 논문수", "1인당 논문수", f"{region_name} 순위", "전국 순위"]

    table = doc.add_table(rows=1 + len(years), cols=len(headers))
    table.style = "Table Grid"

    # 헤더 행
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(hdr_cells[i], "2E75B6")
        run = hdr_cells[i].paragraphs[0].runs[0]
        _set_font(run, size_pt=9, bold=True, color_hex="FFFFFF")

    # 데이터 행
    for row_idx, year in enumerate(years):
        d = hoseo_trend[year]
        vals = [
            str(year),
            f"{d['전임교원수']:,}명",
            f"{d['논문수']:.2f}편",
            f"{d['1인당논문수']:.4f}",
            f"{d['권역순위']}위" if d['권역순위'] else "-",
            f"{d['전국순위']}위",
        ]
        row_cells = table.rows[row_idx + 1].cells
        bg = "D6E4F0" if row_idx % 2 == 0 else "FFFFFF"
        for i, v in enumerate(vals):
            row_cells[i].text = v
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(row_cells[i], bg)
            run = row_cells[i].paragraphs[0].runs[0]
            _set_font(run, size_pt=9)

    doc.add_paragraph()


def _add_compare_group_table(doc: Document, compare_data: list[dict], year: int, univ: str = UNIVERSITY, region_name: str = "충청권"):
    """비교군 5개교 현황 표.

    Args:
        doc: python-docx Document 객체
        compare_data: 비교군 대학별 데이터 리스트
        year: 기준 연도
        univ: 강조 표시할 대학명 (기본값: config의 UNIVERSITY)
        region_name: 권역명 (기본값: "충청권")
    """
    headers = ["대학명", "전임교원수", "논문수", "1인당 논문수", "전국 순위", f"{region_name} 순위"]

    table = doc.add_table(rows=1 + len(compare_data), cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(hdr_cells[i], "2E75B6")
        run = hdr_cells[i].paragraphs[0].runs[0]
        _set_font(run, size_pt=9, bold=True, color_hex="FFFFFF")

    for row_idx, d in enumerate(compare_data):
        is_hoseo = d["학교명"] == univ
        vals = [
            d["학교명"],
            f"{d['전임교원수']:,}명",
            f"{d['논문수']:.2f}편",
            f"{d['1인당논문수']:.4f}",
            f"{d['전국순위']}위",
            f"{d['권역순위']}위" if d['권역순위'] else "-",
        ]
        row_cells = table.rows[row_idx + 1].cells
        bg = "FFF2CC" if is_hoseo else ("D6E4F0" if row_idx % 2 == 0 else "FFFFFF")
        for i, v in enumerate(vals):
            row_cells[i].text = v
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(row_cells[i], bg)
            run = row_cells[i].paragraphs[0].runs[0]
            _set_font(run, size_pt=9, bold=is_hoseo)

    doc.add_paragraph()


def _add_yoy_table(doc: Document, yoy_changes: dict, year: int):
    """전년대비 증감 상위·하위 표."""
    headers = ["구분", "대학명", f"{year-1}년 실적", f"{year}년 실적", "증감률(%)"]

    rows_data = []
    for i, d in enumerate(yoy_changes.get("상위", [])):
        rows_data.append((f"상위 {i+1}", d["학교명"], d["비교연도"], d["기준연도"], d["증감률"]))
    for i, d in enumerate(yoy_changes.get("하위", [])):
        rows_data.append((f"하위 {i+1}", d["학교명"], d["비교연도"], d["기준연도"], d["증감률"]))

    if not rows_data:
        doc.add_paragraph("(전년도 데이터 없음)")
        return

    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(hdr_cells[i], "2E75B6")
        run = hdr_cells[i].paragraphs[0].runs[0]
        _set_font(run, size_pt=9, bold=True, color_hex="FFFFFF")

    for row_idx, (gubun, name, prev, curr, rate) in enumerate(rows_data):
        vals = [gubun, name, f"{prev:.4f}", f"{curr:.4f}", f"{rate:+.1f}%"]
        row_cells = table.rows[row_idx + 1].cells
        bg = "D6E4F0" if row_idx % 2 == 0 else "FFFFFF"
        for i, v in enumerate(vals):
            row_cells[i].text = v
            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(row_cells[i], bg)
            run = row_cells[i].paragraphs[0].runs[0]
            _set_font(run, size_pt=9)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# 메인: 보고서 조립
# ---------------------------------------------------------------------------

def build_report(
    year: int,
    hoseo_trend: dict[int, dict],
    averages: dict[int, dict],
    compare_data: list[dict],
    yoy_changes: dict,
    rank_changes: dict[int, dict],
    charts: dict[str, BytesIO],
    narratives: dict[str, str],
    university: str | None = None,
    region_name: str = "충청권",
) -> BytesIO:
    """
    모든 섹션을 조립하여 Word BytesIO 파일을 반환한다.

    Args:
        year: 기준 연도
        hoseo_trend: 호서대 연도별 추이
        averages: 전국/충청권/비교군 평균
        compare_data: 비교군 5개교 데이터
        yoy_changes: 전년대비 증감 현황
        rank_changes: 순위 변화
        charts: {"trend": BytesIO, "bar": BytesIO, "avg": BytesIO, "rank": BytesIO, "compare": BytesIO}
        narratives: {"trend": str, "comparison": str, "regional": str, "yoy": str}
        university: 보고서에 표시할 대학명. None이면 config의 UNIVERSITY 사용
    """
    # university 파라미터가 없으면 config 기본값 사용
    univ = university or UNIVERSITY
    doc = Document()

    # ----- 페이지 여백 설정 -----
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ========================================================
    # 표지
    # ========================================================
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(f"{univ}")
    _set_font(title_run, size_pt=16, bold=True, color_hex="2E75B6")

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(REPORT_TITLE)
    _set_font(subtitle_run, size_pt=20, bold=True, color_hex="1f4e79")
    subtitle_p.paragraph_format.space_before = Pt(12)

    doc.add_paragraph()
    year_p = doc.add_paragraph()
    year_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    year_run = year_p.add_run(f"기준 연도: {year}년")
    _set_font(year_run, size_pt=13)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = date.today()
    date_run = date_p.add_run(f"생성일: {today.year}년 {today.month}월 {today.day}일")
    _set_font(date_run, size_pt=11)

    doc.add_page_break()

    # ========================================================
    # 섹션 1: 연도별 1인당논문수 추이
    # ========================================================
    _add_heading(doc, "1. 전임교원 1인당 SCI/SCOPUS 논문수 추이")

    if narratives.get("trend"):
        _add_narrative(doc, narratives["trend"])

    _add_trend_table(doc, hoseo_trend, region_name=region_name)

    if charts.get("trend"):
        _add_image(doc, charts["trend"], width_cm=14.0)

    # ========================================================
    # 섹션 2: 전국·충청권·비교군 평균 비교
    # ========================================================
    _add_heading(doc, f"2. 전국·{region_name}·{COMPARE_GROUP_NAME} 평균 비교 ({year}년)")

    if narratives.get("comparison"):
        _add_narrative(doc, narratives["comparison"])

    if charts.get("avg"):
        _add_image(doc, charts["avg"], width_cm=12.0)

    # ========================================================
    # 섹션 3: 충청권 내 비교
    # ========================================================
    _add_heading(doc, f"3. {region_name} 대학 비교 ({year}년)")

    if narratives.get("regional"):
        _add_narrative(doc, narratives["regional"])

    if charts.get("bar"):
        _add_image(doc, charts["bar"], width_cm=15.0)

    if charts.get("rank"):
        _add_image(doc, charts["rank"], width_cm=14.0)

    # ========================================================
    # 섹션 4: 비교군 5개교 현황
    # ========================================================
    _add_heading(doc, f"4. {COMPARE_GROUP_NAME} 현황 ({year}년)")

    _add_compare_group_table(doc, compare_data, year, univ=univ, region_name=region_name)

    if charts.get("compare"):
        _add_image(doc, charts["compare"], width_cm=11.0)

    # ========================================================
    # 섹션 5: 전년대비 증감 현황
    # ========================================================
    _add_heading(doc, f"5. 전년대비 증감 현황 ({year-1}→{year}년)")

    if narratives.get("yoy"):
        _add_narrative(doc, narratives["yoy"])

    _add_yoy_table(doc, yoy_changes, year)

    # ========================================================
    # 저장
    # ========================================================
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
